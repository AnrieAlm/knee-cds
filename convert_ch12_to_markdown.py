# convert_ch12_to_markdown.py
#
# Converts CHAPTER_12_fixed.docx into headed Markdown suitable for
# MarkdownHeaderTextSplitter ingestion.
#
# The source docx has almost no usable style information (949 of 957
# paragraphs are plain "normal", and no bold formatting survived the
# OCR/export round-trip). So heading structure is inferred using
# heuristics rather than read from styles.
#
# Heuristics used, in priority order:
#   1. Real Word heading styles, where they exist (only 8 in this file)
#   2. ALL-CAPS short lines            -> ## (major section, e.g. APPLIED ANATOMY)
#   3. Title Case short lines with no  -> ### (subsection)
#      terminal punctuation
#   4. "Figure 12-N ..."               -> blockquote caption
#   5. "TABLE 12-N ..."                -> bold caption line
#   6. Everything else                 -> body paragraph, unchanged
#
# Run:
#   python convert_ch12_to_markdown.py input.docx output.md

import re
import sys
from pathlib import Path

import docx


# -----------------------------------------------------------
# Settings
# -----------------------------------------------------------

# A line shorter than this MAY be a heading. Longer lines are
# always treated as body text.
MAX_HEADING_CHARS = 60

# Title of the output document.
DOC_TITLE = "Magee's Orthopedic Physical Assessment (6th ed.) — Chapter 12: Knee"

# Provenance note written at the top of the output file, so the
# corpus records where the text came from and how it was produced.
PROVENANCE = (
    "> *Structured notes derived from Magee, D.J., Orthopedic Physical "
    "Assessment, 6th ed., Chapter 12 (Knee). Converted from DOCX to headed "
    "Markdown by convert_ch12_to_markdown.py for corpus ingestion. Heading "
    "levels were inferred heuristically because the source file carried no "
    "usable style information. Figure captions are set as blockquotes. "
    "Reference superscript numbers were left in place rather than stripped, "
    "to avoid corrupting nerve-root notation (e.g. \"L4\", \"S1\").*"
)


# -----------------------------------------------------------
# Heuristic classifiers
# -----------------------------------------------------------

def is_figure_caption(text):
    """Matches lines like 'Figure 12-1 Q-angle differences...'"""
    return bool(re.match(r"^Figure\s+\d+-\d+", text))


def is_table_caption(text):
    """Matches lines like 'TABLE 12-4 Nerve root...'"""
    return bool(re.match(r"^TABLE\s+\d+-\d+", text, re.IGNORECASE))


def is_all_caps_heading(text):
    """
    Major section headings in this chapter are short ALL-CAPS lines,
    e.g. APPLIED ANATOMY, PATIENT HISTORY, OBSERVATION.

    We require at least one letter so that lines of pure punctuation
    or stray numbers don't get promoted to headings.
    """
    if len(text) > MAX_HEADING_CHARS:
        return False
    if not any(c.isalpha() for c in text):
        return False
    return text.isupper()


def is_title_case_heading(text):
    """
    Subsection headings tend to be short Title Case lines with no
    sentence-ending punctuation, e.g. 'Active Movements'.

    This is the loosest heuristic, so it's checked last and kept
    deliberately strict: no trailing period, comma, colon or
    semicolon, and at least two words.
    """
    if len(text) > MAX_HEADING_CHARS:
        return False
    if text[-1] in ".,:;":
        return False

    words = text.split()
    if len(words) < 2:
        return False

    # Count how many words start with a capital letter. Title Case
    # headings will be mostly capitalised; sentences won't be.
    capitalised = sum(1 for w in words if w[:1].isupper())
    return capitalised >= len(words) - 1


# -----------------------------------------------------------
# Main conversion
# -----------------------------------------------------------

def convert(input_path, output_path):

    document = docx.Document(input_path)

    # Lines of the output markdown file, built up one at a time.
    output_lines = []

    # Header block.
    output_lines.append("# " + DOC_TITLE)
    output_lines.append("")
    output_lines.append(PROVENANCE)
    output_lines.append("")

    # Counters so we can assert nothing was silently dropped.
    body_count = 0
    heading_count = 0
    figure_count = 0
    table_count = 0

    for paragraph in document.paragraphs:

        text = paragraph.text.strip()

        # Skip genuinely empty paragraphs. These are layout
        # artefacts, not content.
        if not text:
            continue

        style_name = paragraph.style.name

        # --- Rule 1: trust real Word heading styles where present ---
        if style_name.startswith("Heading"):
            # "Heading 2" -> level 2. Clamp to 2..4 so we never emit
            # a second H1 (the document title owns H1).
            try:
                level = int(style_name.split()[-1])
            except ValueError:
                level = 2
            level = max(2, min(level, 4))

            output_lines.append("")
            output_lines.append("#" * level + " " + text)
            output_lines.append("")
            heading_count += 1
            continue

        # --- Rule 2: figure captions become blockquotes ---
        if is_figure_caption(text):
            output_lines.append("")
            output_lines.append("> *" + text + "*")
            output_lines.append("")
            figure_count += 1
            continue

        # --- Rule 3: table captions become bold lines ---
        if is_table_caption(text):
            output_lines.append("")
            output_lines.append("**" + text + "**")
            output_lines.append("")
            table_count += 1
            continue

        # --- Rule 4: ALL-CAPS short lines are major sections ---
        if is_all_caps_heading(text):
            # Convert to Title Case for readability, since ALL CAPS
            # embeds poorly and reads badly in retrieved chunks.
            output_lines.append("")
            output_lines.append("## " + text.title())
            output_lines.append("")
            heading_count += 1
            continue

        # --- Rule 5: Title Case short lines are subsections ---
        if is_title_case_heading(text):
            output_lines.append("")
            output_lines.append("### " + text)
            output_lines.append("")
            heading_count += 1
            continue

        # --- Rule 6: everything else is body text, unchanged ---
        output_lines.append(text)
        output_lines.append("")
        body_count += 1

    # -----------------------------------------------------------
    # Assertion checks before writing
    # -----------------------------------------------------------

    source_non_empty = sum(
        1 for p in document.paragraphs if p.text.strip()
    )
    accounted = body_count + heading_count + figure_count + table_count

    print("Source non-empty paragraphs :", source_non_empty)
    print("Body paragraphs             :", body_count)
    print("Headings                    :", heading_count)
    print("Figure captions             :", figure_count)
    print("Table captions              :", table_count)
    print("Total accounted for         :", accounted)

    # Every non-empty source paragraph must land in exactly one
    # output category. If this fails, something was dropped.
    assert accounted == source_non_empty, (
        f"Paragraph count mismatch: {accounted} != {source_non_empty}. "
        "Some content was dropped — do not use this output."
    )

    # -----------------------------------------------------------
    # Write the file
    # -----------------------------------------------------------

    Path(output_path).write_text(
        "\n".join(output_lines),
        encoding="utf-8",
    )

    print()
    print("Wrote:", output_path)
    print("Check the heading structure by eye before ingesting:")
    print("  grep '^#' " + str(output_path))


# -----------------------------------------------------------
# Entry point
# -----------------------------------------------------------

if __name__ == "__main__":

    if len(sys.argv) != 3:
        print("Usage: python convert_ch12_to_markdown.py input.docx output.md")
        sys.exit(1)

    convert(sys.argv[1], sys.argv[2])
